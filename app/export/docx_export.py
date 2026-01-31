# -*- coding: utf-8 -*-
"""
Экспорт результатов в Word (.docx)
"""

from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..models.project import Project
from ..models.calculation import CalculationResult


def export_to_docx(project: Project, result: CalculationResult, file_path: str):
    """Экспорт результатов расчёта в Word документ"""
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # === 1. Титульная страница ===
    _add_title_page(doc, project)

    # === 2. Общие сведения о ПС ===
    doc.add_heading('1. Общие сведения о программном средстве', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    rows_data = [
        ('Название ПС', project.name),
        ('Описание', project.description or '-'),
        ('Фонд рабочего времени', f'{project.work_fund} дней/месяц'),
        ('Тип ограничения', 'По продолжительности' if project.constraint_type == 'duration' else 'По численности'),
        ('Значение ограничения', f'{project.constraint_value} {"месяцев" if project.constraint_type == "duration" else "человек"}'),
    ]

    for i, (name, value) in enumerate(rows_data):
        table.rows[i].cells[0].text = name
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph()

    # === 3. Перечень компонентов ===
    doc.add_heading('2. Перечень компонентов', level=1)

    for i, comp in enumerate(project.components, 1):
        doc.add_heading(f'2.{i}. {comp.name}', level=2)
        if comp.description:
            doc.add_paragraph(comp.description)

        if comp.functions:
            table = doc.add_table(rows=len(comp.functions) + 1, cols=5)
            table.style = 'Table Grid'

            # Заголовки
            headers = ['№', 'Функция', 'Vi', 'ri', 'ki']
            for j, h in enumerate(headers):
                table.rows[0].cells[j].text = h
                table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

            # Данные
            for k, func in enumerate(comp.functions):
                table.rows[k + 1].cells[0].text = func.function_id
                table.rows[k + 1].cells[1].text = func.function_name[:50]
                table.rows[k + 1].cells[2].text = str(func.volume)
                table.rows[k + 1].cells[3].text = str(func.reuse_count)
                table.rows[k + 1].cells[4].text = f'{func.reuse_coefficient:.2f}'

    doc.add_paragraph()

    # === 4. Каталог функций ===
    doc.add_heading('3. Детальный каталог функций', level=1)

    table = doc.add_table(rows=len(result.functions_results) + 1, cols=7)
    table.style = 'Table Grid'

    headers = ['Компонент', 'Функция', 'Vi', 'ri', 'ki', 'Vm', 'Vk']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    for k, fr in enumerate(result.functions_results):
        table.rows[k + 1].cells[0].text = fr.component_name[:15]
        table.rows[k + 1].cells[1].text = fr.function_name[:30]
        table.rows[k + 1].cells[2].text = str(fr.volume_base)
        table.rows[k + 1].cells[3].text = str(fr.reuse_count)
        table.rows[k + 1].cells[4].text = f'{fr.reuse_coefficient:.2f}'
        table.rows[k + 1].cells[5].text = f'{fr.volume_corrected:.0f}'
        table.rows[k + 1].cells[6].text = f'{fr.volume_adjusted:.0f}'

    doc.add_paragraph()
    doc.add_paragraph(f'Общий объём ПС: V = {result.total_volume:.2f} строк условного кода')

    # === 5. Коэффициенты уровня расчёта ===
    doc.add_heading('4. Коэффициенты уровня расчёта', level=1)

    coeffs = project.coefficients
    coef_data = [
        ('K_n (степень новизны)', coeffs.novelty, result.k_n),
        ('K_nad (надёжность)', coeffs.reliability, result.k_nad),
        ('K_proizv (производительность)', coeffs.performance, result.k_proizv),
        ('K_dokum (документация)', coeffs.documentation, result.k_dokum),
        ('K_teh (технологии)', f'K_str × Π(K_t)', result.k_teh),
        ('K_or (опыт разработки)', coeffs.dev_experience, result.k_or),
    ]

    table = doc.add_table(rows=len(coef_data) + 1, cols=3)
    table.style = 'Table Grid'

    headers = ['Коэффициент', 'Значение', 'Числовое значение']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    for k, (name, value, num) in enumerate(coef_data):
        table.rows[k + 1].cells[0].text = name
        table.rows[k + 1].cells[1].text = str(value)[:50]
        table.rows[k + 1].cells[2].text = f'{num:.4f}'

    doc.add_paragraph()

    # === 6. Расчёт базовой трудоёмкости ===
    doc.add_heading('5. Расчёт базовой трудоёмкости', level=1)

    doc.add_paragraph('Формула расчёта базовой трудоёмкости:')
    doc.add_paragraph('T_baz = A × V^C × K_n × K_nad × K_proizv × K_dokum × K_teh × K_or')

    formula_text = (
        f'T_baz = {result.A} × {result.total_volume:.2f}^{result.C} × '
        f'{result.k_n} × {result.k_nad} × {result.k_proizv} × '
        f'{result.k_dokum} × {result.k_teh:.4f} × {result.k_or}'
    )
    doc.add_paragraph(formula_text)

    p = doc.add_paragraph()
    run = p.add_run(f'T_baz = {result.base_labor:.2f} чел.-дн.')
    run.bold = True

    doc.add_paragraph()

    # === 7. Таблица подпроцессов ===
    doc.add_heading('6. Трудоёмкость подпроцессов', level=1)

    table = doc.add_table(rows=len(result.subprocess_results) + 2, cols=5)
    table.style = 'Table Grid'

    headers = ['Подпроцесс', 'Коэфф. A', 'Трудоёмкость [чел.-дн.]', 'Численность [чел.]', 'Срок [мес.]']
    for j, h in enumerate(headers):
        table.rows[0].cells[j].text = h
        table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

    for k, sp in enumerate(result.subprocess_results):
        table.rows[k + 1].cells[0].text = sp.name
        table.rows[k + 1].cells[1].text = f'{sp.base_coefficient}'
        table.rows[k + 1].cells[2].text = f'{sp.labor:.2f}'
        table.rows[k + 1].cells[3].text = f'{sp.staff:.2f}'
        table.rows[k + 1].cells[4].text = f'{sp.duration:.2f}'

    # Итого
    last_row = len(result.subprocess_results) + 1
    table.rows[last_row].cells[0].text = 'ИТОГО'
    table.rows[last_row].cells[1].text = '1.00'
    table.rows[last_row].cells[2].text = f'{result.total_labor:.2f}'
    table.rows[last_row].cells[3].text = f'{result.average_staff:.2f}'
    table.rows[last_row].cells[4].text = f'{result.total_duration:.2f}'

    for cell in table.rows[last_row].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # === 8. Выводы ===
    doc.add_heading('7. Итоговые показатели', level=1)

    conclusions = [
        f'Трудоёмкость разработки: T_razr = {result.total_labor:.2f} чел.-дн.',
        f'Коэффициент сокращения сроков: K_sr_srok = {result.k_sr_srok}',
        f'Итоговая трудоёмкость: T_srok = {result.final_labor:.2f} чел.-дн.',
        f'Общий срок разработки: {result.total_duration:.2f} месяцев',
        f'Средняя численность: {result.average_staff:.2f} человек',
    ]

    for c in conclusions:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        f'Расчёт выполнен {datetime.now().strftime("%d.%m.%Y %H:%M")} '
        f'с использованием методики СПбГУТ им. проф. М.А. Бонч-Бруевича.'
    )

    # Сохранение
    doc.save(file_path)


def _add_title_page(doc: Document, project: Project):
    """Добавить титульную страницу"""
    # Заголовок
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('РАСЧЁТ ТРУДОЁМКОСТИ РАЗРАБОТКИ')
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ПРОГРАММНОГО СРЕДСТВА')
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()
    doc.add_paragraph()

    # Название ПС
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(project.name)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph()
    doc.add_paragraph()

    # Методика
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('по методике СПбГУТ им. проф. М.А. Бонч-Бруевича')

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Дата
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(datetime.now().strftime('%d.%m.%Y'))

    # Разрыв страницы
    doc.add_page_break()
